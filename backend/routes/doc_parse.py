"""
Document Parser Route
=====================
Two endpoints:

POST /api/documents/parse
  Extracts text + uses LLM to produce a structured requirement string (fills textarea).

POST /api/documents/generate
  Full pipeline: parse file → extract structured tags → LLM graph extraction →
  assemble IEC 61131-3 ST code → IEC validate → return deployable code.
  No manual description needed — the file IS the spec.
"""

from __future__ import annotations
import io
import json
import re

from fastapi import APIRouter, UploadFile, File, Form, HTTPException, Request
from backend.openai_client import safe_chat_completion, DEFAULT_MODEL

router = APIRouter()


# ── Text extractors ────────────────────────────────────────────────────────────

def _extract_pdf(data: bytes) -> str:
    try:
        import pdfplumber
    except ImportError:
        raise HTTPException(500, "pdfplumber not installed on server")

    parts: list[str] = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            txt = page.extract_text()
            if txt:
                parts.append(txt.strip())
            for table in (page.extract_tables() or []):
                for row in table:
                    row_str = " | ".join(str(c).strip() for c in row if c and str(c).strip())
                    if row_str:
                        parts.append(row_str)
    return "\n".join(parts)


def _extract_docx(data: bytes) -> str:
    try:
        from docx import Document
    except ImportError:
        raise HTTPException(500, "python-docx not installed on server")

    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_xlsx(data: bytes) -> str:
    try:
        import openpyxl
    except ImportError:
        raise HTTPException(500, "openpyxl not installed on server")

    wb = openpyxl.load_workbook(io.BytesIO(data), data_only=True)
    parts: list[str] = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        parts.append(f"=== Sheet: {sheet_name} ===")
        for row in ws.iter_rows(values_only=True):
            vals = [str(v).strip() for v in row if v is not None and str(v).strip()]
            if vals:
                parts.append(" | ".join(vals))
    return "\n".join(parts)


def _extract_hmi_json(data: bytes) -> str:
    """Parse an exported HMI JSON and summarise widget/tag info."""
    try:
        obj = json.loads(data.decode("utf-8", errors="replace"))
    except Exception:
        raise HTTPException(400, "Invalid JSON file")

    lines: list[str] = []
    if isinstance(obj, dict):
        if "widgets" in obj:
            lines.append("HMI Widgets:")
            for w in obj["widgets"]:
                tag  = w.get("tag") or w.get("id") or ""
                kind = w.get("type") or w.get("widget_type") or ""
                lbl  = w.get("label") or w.get("title") or ""
                lines.append(f"  - {kind}: {lbl} (tag={tag})")
        if "title" in obj:
            lines.insert(0, f"HMI Screen: {obj['title']}")
    elif isinstance(obj, list):
        for item in obj:
            if isinstance(item, dict):
                lines.append(json.dumps(item))
    return "\n".join(lines) if lines else json.dumps(obj, indent=2)[:4000]


# ── LLM extraction ─────────────────────────────────────────────────────────────

_SYSTEM = """You are a senior industrial automation engineer reading a document to generate IEC 61131-3 PLC code.

Your job: extract ALL control requirements from the document and write a single clear, structured description.

OUTPUT FORMAT — use exactly this structure:

[MACHINE / PROCESS]
<What the machine does — one concise paragraph. Name the process, what it controls, end product or purpose.>

[OPERATING STATES]
List every operating state the machine goes through, in order:
  Idle → <State2> → <State3> → ... → Fault
Include: what triggers each state change, what outputs are active in each state.

[INPUTS — Digital]
- I_<TagName> (NO|NC) — <description, physical device, location if given>
Always include: I_Start, I_Stop, I_EStop (Emergency Stop, NC), I_Reset.
Add every sensor, limit switch, pressure switch, level switch, proximity sensor mentioned.

[INPUTS — Analog] (omit section if none)
- AI_<TagName> (4–20mA|0–10V) — <description, range, engineering units>

[OUTPUTS — Digital]
- Q_<TagName> — <actuator or device driven, e.g. motor contactor, solenoid valve>
Always include Q_Fault (fault indicator).

[OUTPUTS — Analog] (omit section if none)
- AO_<TagName> — <description, range>

[TIMERS]
List every timed operation:
- <TimerName>: <duration>, active in <State>, purpose: <what it delays>
If no timers mentioned, write: None specified.

[INTERLOCKS & SAFETY]
- E-Stop: <behavior — latching or non-latching, which outputs it cuts>
- Thermal overload: <behavior>
- Any other safety interlocks
- Fault reset: <how operator resets faults>

[ALARMS]
- <alarm condition> → <action taken>

[SEQUENCE / CONTROL LOGIC]
Step-by-step description of the control sequence. Be specific:
  1. In Idle: waiting for Start signal...
  2. On Start pressed: ...
  3. After <condition>: transition to <next state>...

RULES:
- Use industrial tag naming: I_ for inputs, Q_ for outputs, AI_/AO_ for analog.
- Always specify NO (normally open) or NC (normally closed) for digital inputs.
- If information is missing from the document, write "Not specified" — do NOT invent data.
- Be concise but complete. This text goes directly into an industrial PLC code generator."""


def _llm_extract(raw_text: str) -> str:
    truncated = raw_text[:8000]  # ~2k tokens of document content
    try:
        resp = safe_chat_completion(
            model=DEFAULT_MODEL,
            messages=[
                {"role": "system", "content": _SYSTEM},
                {"role": "user",   "content": f"Document content to extract PLC requirements from:\n\n{truncated}"},
            ],
            temperature=0.0,
            max_tokens=2000,
        )
        return resp.choices[0].message.content.strip()
    except Exception as e:
        raise HTTPException(500, f"LLM extraction failed: {e}")


# ── Route ──────────────────────────────────────────────────────────────────────

@router.post("/api/documents/parse")
async def parse_document(file: UploadFile = File(...)):
    """
    Parse an uploaded document and return a PLC requirement string.
    Supported: PDF, DOCX, XLSX/XLS, JSON (HMI export)
    """
    filename  = (file.filename or "").lower()
    data      = await file.read()

    if not data:
        raise HTTPException(400, "Empty file uploaded")

    # Determine file type and extract raw text
    if filename.endswith(".pdf"):
        raw  = _extract_pdf(data)
        ftype = "PDF"
    elif filename.endswith(".docx") or filename.endswith(".doc"):
        raw  = _extract_docx(data)
        ftype = "DOCX"
    elif filename.endswith(".xlsx") or filename.endswith(".xls"):
        raw  = _extract_xlsx(data)
        ftype = "XLSX"
    elif filename.endswith(".json"):
        raw  = _extract_hmi_json(data)
        ftype = "JSON"
    else:
        raise HTTPException(400, f"Unsupported file type: {filename}. Use PDF, DOCX, XLSX, or JSON.")

    if not raw.strip():
        raise HTTPException(422, "Could not extract any text from the document. Check the file is not image-only or empty.")

    # Use LLM to produce a clean requirement description
    requirement = _llm_extract(raw)

    return {
        "requirement": requirement,
        "file_type":   ftype,
        "raw_chars":   len(raw),
    }


# ════════════════════════════════════════════════════════════════════════════
# DIRECT GENERATION PIPELINE  —  file → graph → IEC code
# ════════════════════════════════════════════════════════════════════════════

# Common column name variants for smart Excel tag parsing
_TAG_HEADERS   = {"tag", "tag name", "name", "signal name", "i/o name", "point name",
                  "variable", "tag_name", "tagname", "io name", "field tag", "instrument tag",
                  "instrument", "device", "label"}
_DIR_HEADERS   = {"direction", "type", "i/o", "io", "input/output", "io type", "di/do/ai/ao",
                  "signal direction", "i/o type", "in/out", "io_type", "field type",
                  "channel type", "signal class"}
_DTYPE_HEADERS = {"signal", "signal type", "no/nc", "contact", "dtype", "data type",
                  "contact type", "wiring", "nc/no", "normally open", "normally closed"}
_ADDR_HEADERS  = {"address", "addr", "plc address", "hw address", "channel", "module",
                  "hardware address", "i/o address", "byte.bit", "rack", "slot",
                  "plc tag", "memory address", "bit address", "physical address"}
_DESC_HEADERS  = {"description", "desc", "notes", "function", "comment", "purpose",
                  "remarks", "detail", "explanation", "note", "function description"}


def _parse_excel_tags(data: bytes) -> list[dict]:
    """
    Detect an I/O tag table in an Excel workbook.
    Returns a list of dicts with keys: name, direction, signal_type, address, desc.
    Falls back to empty list if no tag table detected.
    """
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(data), data_only=True)
    tags: list[dict] = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = [
            [str(c.value).strip() if c.value is not None else "" for c in row]
            for row in ws.iter_rows()
        ]
        if not rows:
            continue

        # Find header row — first row containing a tag-like column
        header_idx = -1
        col_map: dict[str, int] = {}
        for ri, row in enumerate(rows[:30]):  # search first 30 rows (handles title/logo rows)
            lrow = [c.lower().strip() for c in row]
            found = {}
            for ci, cell in enumerate(lrow):
                # Exact match
                if cell in _TAG_HEADERS and "tag" not in found:
                    found["tag"] = ci
                if cell in _DIR_HEADERS and "dir" not in found:
                    found["dir"] = ci
                if cell in _DTYPE_HEADERS and "dtype" not in found:
                    found["dtype"] = ci
                if cell in _ADDR_HEADERS and "addr" not in found:
                    found["addr"] = ci
                if cell in _DESC_HEADERS and "desc" not in found:
                    found["desc"] = ci
                # Partial / contains match (handles "Tag Name / Signal", "I/O Address" etc.)
                if "tag" not in found and any(h in cell for h in ("tag", "signal name", "point")):
                    found["tag"] = ci
                if "dir" not in found and any(h in cell for h in ("direction", "input", "output", "i/o")):
                    found["dir"] = ci
                if "dtype" not in found and any(h in cell for h in ("no/nc", "nc/no", "contact type", "wiring")):
                    found["dtype"] = ci
                if "addr" not in found and any(h in cell for h in ("address", "addr", "channel")):
                    found["addr"] = ci
                if "desc" not in found and any(h in cell for h in ("desc", "note", "remark", "function")):
                    found["desc"] = ci
            if "tag" in found:
                header_idx = ri
                col_map = found
                break

        if header_idx == -1:
            continue  # no header found in this sheet

        for row in rows[header_idx + 1:]:
            name = row[col_map["tag"]] if "tag" in col_map and col_map["tag"] < len(row) else ""
            if not name or name.lower() in ("", "none", "tag name", "name"):
                continue

            direction   = row[col_map["dir"]]   if "dir"   in col_map and col_map["dir"]   < len(row) else ""
            signal_type = row[col_map["dtype"]]  if "dtype" in col_map and col_map["dtype"] < len(row) else ""
            address     = row[col_map["addr"]]   if "addr"  in col_map and col_map["addr"]  < len(row) else ""
            desc        = row[col_map["desc"]]   if "desc"  in col_map and col_map["desc"]  < len(row) else ""

            # Normalise direction
            d_low = direction.lower()
            if any(k in d_low for k in ("di", "input", "in")):
                direction = "INPUT"
            elif any(k in d_low for k in ("do", "output", "out")):
                direction = "OUTPUT"
            elif any(k in d_low for k in ("ai",)):
                direction = "ANALOG_IN"
            elif any(k in d_low for k in ("ao",)):
                direction = "ANALOG_OUT"

            tags.append({
                "name":        name,
                "direction":   direction,
                "signal_type": signal_type.upper() if signal_type else "",
                "address":     address,
                "desc":        desc,
            })

        if tags:
            break  # found tags in this sheet; stop searching

    return tags


# ── Graph-extract prompt for direct file-based generation ─────────────────────

_GRAPH_SYSTEM_DOC = (
    "Return JSON only. No markdown, no explanation, no text before or after the JSON."
)

_GRAPH_EXAMPLE_DOC = """\
EXAMPLE — incinerator system:
{
  "states": ["Idle", "Preheat", "Burning", "Cooldown", "Fault"],
  "inputs": ["DI_START","DI_STOP","DI_DOOR","DI_TRAY","DI_WATER_PUMP","AI_BURN_TEMP","AI_ROOM_TEMP"],
  "outputs": ["DO_IGNITION","DO_FAN","DO_BURNER","Q_Fault"],
  "timers": [
    {"name":"T_Preheat","preset":"T#600S","active_in":["Preheat"]},
    {"name":"T_Cooldown","preset":"T#300S","active_in":["Cooldown"]}
  ],
  "outputs_active_in": {
    "Preheat":  ["DO_IGNITION","DO_FAN"],
    "Burning":  ["DO_IGNITION","DO_FAN","DO_BURNER"],
    "Cooldown": ["DO_FAN"],
    "Fault":    ["Q_Fault"]
  },
  "transitions": {
    "Idle":     [{"cond":"DI_START AND DI_DOOR AND DI_TRAY AND DI_WATER_PUMP","next":"Preheat"}],
    "Preheat":  [{"cond":"AI_BURN_TEMP >= 400","next":"Burning"},
                 {"cond":"DI_STOP","next":"Idle"}],
    "Burning":  [{"cond":"DI_STOP","next":"Cooldown"}],
    "Cooldown": [{"cond":"T_Cooldown.Q AND AI_BURN_TEMP < 80","next":"Idle"}],
    "Fault":    [{"cond":"DI_RESET AND NOT DI_DOOR_FAULT","next":"Idle"}]
  },
  "fault_sources": [
    {"var":"AI_BURN_TEMP >= 600","code":10,"desc":"High Temp"},
    {"var":"NOT DI_DOOR","code":40,"desc":"Door Open"},
    {"var":"NOT DI_TRAY","code":41,"desc":"Tray Missing"}
  ]
}"""


def _build_doc_prompt(raw_text: str, tags: list[dict], program_name: str) -> str:
    """Build the LLM prompt that feeds into extract_plc_graph."""
    lines = [_GRAPH_EXAMPLE_DOC, ""]

    if tags:
        lines.append("=== EXACT I/O TAGS FROM UPLOADED FILE — USE THESE NAMES VERBATIM ===")
        inputs  = [t for t in tags if "INPUT"  in t["direction"] or t["name"].upper().startswith(("DI_","I_","AI_"))]
        outputs = [t for t in tags if "OUTPUT" in t["direction"] or t["name"].upper().startswith(("DO_","Q_","AO_"))]
        if inputs:
            lines.append("INPUTS:")
            for t in inputs:
                lines.append(f"  {t['name']} ({t.get('signal_type','') or t['direction']}) — {t['desc'] or 'no description'}")
        if outputs:
            lines.append("OUTPUTS:")
            for t in outputs:
                lines.append(f"  {t['name']} — {t['desc'] or 'no description'}")
        lines.append("=== END TAGS ===\n")

    lines.append("=== DOCUMENT CONTENT ===")
    lines.append(raw_text[:6000])
    lines.append("=== END DOCUMENT ===\n")

    lines.append(
        f"Now build a complete state machine graph for program '{program_name}'.\n"
        "MANDATORY:\n"
        "1. Use the EXACT tag names listed above — do NOT rename or simplify them.\n"
        "2. 'states': first = 'Idle', last = 'Fault'. Add every process state from the document.\n"
        "3. 'inputs': list all I/O names from the tag table above.\n"
        "4. 'outputs': list all output tag names + Q_Fault.\n"
        "5. 'outputs_active_in': every non-Idle, non-Fault state must activate at least one output.\n"
        "6. 'transitions': every state needs at least one exit. Use exact tag names in conditions.\n"
        "7. 'fault_sources': list every fault condition with code and description from the document.\n"
        "8. 'timers': include every timed operation with realistic preset values.\n"
        "Return ONLY the JSON — no other text."
    )
    return "\n".join(lines)


# ── Direct generation endpoint ────────────────────────────────────────────────

@router.post("/api/documents/generate")
async def generate_from_document(
    request:     Request,
    file:         UploadFile = File(...),
    program_name: str        = Form("MyPLCProgram"),
    brand:        str        = Form("siemens"),
    user_email:   str        = Form(""),        # backup: also sent in form body
    description:  str        = Form(""),        # optional user requirements text
):
    """
    Full pipeline: upload a spec file → extract tags → graph extraction →
    assemble IEC 61131-3 ST → IEC validate → return production-ready code.

    Form fields:
      file         — PDF / DOCX / XLSX / JSON
      program_name — PLC program name (default: MyPLCProgram)
      brand        — siemens | codesys | allenbradley | schneider
      user_email   — optional backup for X-User-Email header
    """
    # Email: prefer X-User-Email header; fall back to form field
    email    = request.headers.get("X-User-Email") or user_email or None
    filename = (file.filename or "").lower()
    print(f"[DOC-GEN] START  file={filename}  email={email!r}  brand={brand}")
    data     = await file.read()

    if not data:
        raise HTTPException(400, "Empty file uploaded")

    # ── 1. Parse file to text + structured tags ───────────────────────────────
    tags: list[dict] = []
    if filename.endswith(".pdf"):
        raw   = _extract_pdf(data)
        ftype = "PDF"
    elif filename.endswith(".docx") or filename.endswith(".doc"):
        raw   = _extract_docx(data)
        ftype = "DOCX"
    elif filename.endswith(".xlsx") or filename.endswith(".xls"):
        raw   = _extract_xlsx(data)
        tags  = _parse_excel_tags(data)   # also try structured parse
        ftype = "XLSX"
    elif filename.endswith(".json"):
        raw   = _extract_hmi_json(data)
        ftype = "JSON"
    else:
        raise HTTPException(400, f"Unsupported file type. Use PDF, DOCX, XLSX, or JSON.")

    if not raw.strip() and not tags:
        raise HTTPException(422, "Could not extract content from file.")

    # Merge user requirements with file content (user text takes priority)
    if description and description.strip():
        raw = f"=== USER REQUIREMENTS ===\n{description.strip()}\n\n=== FILE CONTENT ===\n{raw}"
        print(f"[DOC-GEN] User requirements merged ({len(description)} chars)")

    # ── 2. Pre-flight quota check ─────────────────────────────────────────────
    if email:
        from backend.token_manager import check_and_update_tokens
        chk = check_and_update_tokens(email, 0)
        if chk.get("blocked"):
            raise HTTPException(403, "Token limit reached. Please upgrade to continue.")

    from backend.enhanced_intelligent_generator import (
        generate_st_direct,
        extract_plc_graph, _validate_and_fix_graph,
        _extract_timers_from_prompt, _wire_timers_into_transitions,
        assemble_st, build_tag_table, _auto_assign_addresses,
    )
    from backend.iec_final_fixer import final_iec_fix
    from backend.industrial_iec_validator import IndustrialIECValidator
    from backend.engine.vendor_profile import VendorProfileInjector

    # ── 3. PRIMARY: Direct LLM generation — full context, no graph compression ─
    # The LLM writes complete IEC code using all tag data + process description.
    # This produces the same quality as asking Claude/OpenAI directly.
    code = ""
    tokens = 0
    generation_path = "unknown"

    try:
        print(f"[DOC-GEN] Attempting direct LLM generation (tags={len(tags)})...")
        code, tokens = generate_st_direct(raw[:8000], tags, program_name, brand)
        if code and len(code) > 200 and "FUNCTION_BLOCK" in code.upper():
            generation_path = "direct_llm"
            print(f"[DOC-GEN] Direct LLM generation succeeded: {len(code)} chars, {tokens} tokens")
        else:
            print(f"[DOC-GEN] Direct LLM output too short or missing structure — falling back")
            code = ""
    except Exception as e:
        print(f"[DOC-GEN] Direct LLM generation failed ({e}), falling back to graph pipeline")

    # ── 4. FALLBACK: Graph extraction → Python assembly (original pipeline) ────
    if not code:
        try:
            print(f"[DOC-GEN] Attempting graph-based fallback...")
            # Build prompt WITHOUT wrapping in extract_plc_graph's template again
            # (fixes the double-prompt wrap that confused the LLM)
            rich_prompt = _build_doc_prompt(raw, tags, program_name)
            graph, tokens = extract_plc_graph(rich_prompt, program_name)

            graph["timers"] = _extract_timers_from_prompt(rich_prompt, graph)
            _wire_timers_into_transitions(graph)
            all_conds = " ".join(
                t.get("cond", "")
                for tlist in graph.get("transitions", {}).values()
                for t in tlist
            )
            graph["timers"] = [
                tmr for tmr in graph["timers"]
                if f"{tmr['name']}.Q" in all_conds or f"{tmr['name']}.ET" in all_conds
            ]

            # Merge any Excel tags the LLM missed
            if tags:
                existing_inputs  = set(graph.get("inputs",  []))
                existing_outputs = set(graph.get("outputs", []))
                for t in tags:
                    nm = t["name"]
                    if "INPUT" in t.get("direction","") or nm.upper().startswith(("DI_","I_","AI_")):
                        if nm not in existing_inputs:
                            graph.setdefault("inputs", []).append(nm)
                    elif "OUTPUT" in t.get("direction","") or nm.upper().startswith(("DO_","Q_","AO_")):
                        if nm not in existing_outputs:
                            graph.setdefault("outputs", []).append(nm)

            excel_addresses = {t["name"]: t["address"] for t in tags if t.get("address")} if tags else {}
            code = assemble_st(graph, program_name, vendor=brand, description=rich_prompt,
                               addresses_override=excel_addresses if excel_addresses else None)
            generation_path = "graph_fallback"
            print(f"[DOC-GEN] Graph fallback succeeded: {len(code)} chars")
        except Exception as e2:
            raise HTTPException(500, f"Both generation paths failed. Last error: {e2}")

    # ── 5. Final fix + IEC validation ─────────────────────────────────────────
    code = final_iec_fix(code, program_name)
    code = VendorProfileInjector.apply(code, brand, program_name)
    iec  = IndustrialIECValidator.validate(code)
    print(f"[DOC-GEN] IEC validation: valid={iec['valid']} errors={len(iec['errors'])} "
          f"warnings={len(iec['warnings'])} path={generation_path}")

    # ── 8. Build tag table ────────────────────────────────────────────────────
    if tags:
        # Use tags exactly as provided by the Excel file
        def _tag_type(t: dict) -> str:
            nm = t["name"].upper()
            if nm.startswith("AI_") or t["direction"] == "ANALOG_IN":
                return "INT"   # 4-20mA/TC analog — TAG-4/5 compliance
            if nm.startswith("AO_") or t["direction"] == "ANALOG_OUT":
                return "INT"
            return "BOOL"

        tag_table = [
            {
                "tag":         t["name"],
                "address":     t.get("address", ""),
                "type":        _tag_type(t),
                "direction":   t["direction"].replace("ANALOG_IN","INPUT").replace("ANALOG_OUT","OUTPUT"),
                "signal_type": t.get("signal_type", "4-20mA" if t["name"].upper().startswith("AI_") else "NO"),
                "voltage":     "mA" if t["name"].upper().startswith(("AI_","AO_")) else "24VDC",
                "terminal":    t.get("address",""),
                "description": t.get("desc",""),
            }
            for t in tags
        ]
    else:
        addresses = _auto_assign_addresses(
            list(graph.get("inputs",[])), list(graph.get("outputs",[])), brand
        )
        tag_table = build_tag_table(
            list(graph.get("inputs",[])), list(graph.get("outputs",[])), addresses
        )

    # ── 9. Track tokens ───────────────────────────────────────────────────────
    if email and tokens > 0:
        from backend.token_manager import check_and_update_tokens
        result = check_and_update_tokens(email, tokens, "/api/documents/generate")
        print(f"[DOC-GEN] TRACKED  email={email!r}  tokens={tokens}  "
              f"new_total={result.get('tokens_used')}  remaining={result.get('remaining')}")
    else:
        print(f"[DOC-GEN] SKIP TRACK  email={email!r}  tokens={tokens}  "
              f"(email={'missing' if not email else 'ok'}, tokens={'zero - LLM fallback' if tokens == 0 else 'ok'})")

    print(
        f"[DOC-GEN] DONE  {filename} → {ftype} | tags={len(tags)} | tokens={tokens} | "
        f"iec_valid={iec['valid']} | code={len(code)} chars"
    )

    return {
        "code":        code,
        "iec_valid":   iec["valid"],
        "iec_errors":  iec["errors"],
        "iec_warnings":iec["warnings"],
        "domain":      f"doc_{ftype.lower()}_{generation_path}",
        "tokens_used": tokens,
        "tag_list":    tag_table,
        "file_type":   ftype,
        "tags_found":  len(tags),
        "generation_path": generation_path,
        "clarification_needed": False,
    }
